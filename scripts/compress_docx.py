#!/usr/bin/env python3
"""
DOCX compression with hybrid labeled contact format for maximum ATS compatibility
Fixes: No double bullets, proper email with @, Professional Summary
"""
import os
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ python-docx not installed")
    sys.exit(1)

def add_horizontal_line(paragraph, color_rgb=(70, 130, 180)):
    """Add blue underline to section headers"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '8')
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), f'{color_rgb[0]:02x}{color_rgb[1]:02x}{color_rgb[2]:02x}')
    borders.append(bottom_border)
    pPr.append(borders)

def is_bullet_point(text):
    """Detect if a paragraph should be a bullet point"""
    # Skip if already has bullet symbol
    if text.strip().startswith('•'):
        return False
        
    key_skills = [
        "Customer service", "POS system", "Merchandising",
        "Team collaboration", "Technology,", "Loss prevention",
        "Effective communication"
    ]
    
    action_verbs = [
        "Delivered", "Processed", "Handled", "Set up",
        "Trained", "Specialized", "Maintained", 
        "Collaborated", "Applied", "Partnered"
    ]
    
    if any(skill in text for skill in key_skills):
        return True
    
    if any(text.strip().startswith(verb) for verb in action_verbs):
        return True
    
    if text.startswith("Passionate about") or text.startswith("Available for"):
        return True
        
    return False

def add_bullet_symbol(paragraph):
    """Add bullet symbol (•) only if not already present"""
    if paragraph.runs and paragraph.text.strip():
        full_text = paragraph.text.strip()
        
        # Don't add if already has bullet
        if full_text.startswith('•'):
            return
            
        paragraph.clear()
        bullet_text = f"• {full_text}"
        run = paragraph.add_run(bullet_text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)

def compress_docx_formatting(docx_path):
    """Apply complete formatting with hybrid labeled contact for ATS"""
    
    if not os.path.exists(docx_path):
        print(f"❌ File not found: {docx_path}")
        return False
    
    print(f"🔧 Applying ATS-optimized formatting: {docx_path}")
    
    try:
        doc = Document(docx_path)
        
        # Set margins (1.2cm)
        for section in doc.sections:
            section.top_margin = Cm(1.2)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.2)
            section.right_margin = Cm(1.2)
        
        paras_to_remove = []
        
        section_headers = [
            "About Me", "Professional Summary", "Key Skills", "Retail Experience",
            "Other Professional Experience", "Education",
            "Additional Information"
        ]
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # Remove ALL indentation
            para.paragraph_format.left_indent = Pt(0)
            para.paragraph_format.right_indent = Pt(0)
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            
            if not text:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                continue
            
            # HEADER: Remove "Resume - Aaron De Vries" title line
            if text == "Resume - Aaron De Vries":
                paras_to_remove.append(para)
                continue
            
            # HEADER: Main name
            if text == "Aaron De Vries":
                para.clear()
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
                name_run = para.add_run("Aaron De Vries")
                name_run.font.name = "Times New Roman"
                name_run.font.size = Pt(16)
                name_run.font.bold = True
                name_run.font.color.rgb = RGBColor(0, 0, 0)
                continue
            
            # HEADER: Contact block - HYBRID LABELED FORMAT for ATS
            if any(contact in text.lower() for contact in ["homebush", "0400", "@gmail", "@protonmail", "linkedin.com/in"]):
                para.clear()
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(3)
                para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
                # Line 1: Location and Phone
                contact_line1 = "Location: Homebush, NSW  |  Phone: 0400 375 308"
                line1_run = para.add_run(contact_line1)
                line1_run.font.name = "Times New Roman"
                line1_run.font.size = Pt(9)
                line1_run.font.color.rgb = RGBColor(70, 130, 180)
                
                # Add line break
                para.add_run("\n")
                
                # Line 2: Email and LinkedIn (with proper @)
                contact_line2 = "Email: aaron.jp.devries@gmail.com  |  LinkedIn: linkedin.com/in/aarondevriesdev"
                line2_run = para.add_run(contact_line2)
                line2_run.font.name = "Times New Roman"
                line2_run.font.size = Pt(9)
                line2_run.font.color.rgb = RGBColor(70, 130, 180)
                
                para.paragraph_format.space_after = Pt(10)
                continue
            
            # Section headers with blue underlines
            if any(header in text for header in section_headers):
                para.paragraph_format.space_before = Pt(10)
                para.paragraph_format.space_after = Pt(4)
                
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                
                add_horizontal_line(para)
                continue
            
            # BULLET POINTS - only add if not already there
            if is_bullet_point(text):
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.left_indent = Cm(0.5)
                add_bullet_symbol(para)
                continue
            
            # Job titles (bold)
            if any(title in text for title in [
                "Retail Team Member", "Senior Software Engineer",
                "Application Web Developer", "Application & Web Developer",
                "Bachelor of Computer Science"
            ]):
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(1)
                
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                    run.font.bold = True
                continue
            
            # Company/dates lines
            if any(company in text for company in [
                "Dick Smith", "Big W", "Avaloq", "Bond International", "University of Technology Sydney"
            ]) or re.search(r'\d{4}', text):
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(3)
                
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                continue
            
            # Regular content
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(2)
            
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Remove marked paragraphs
        for para in paras_to_remove:
            p = para._element
            p.getparent().remove(p)
        
        # Document-wide defaults
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.0
        
        doc.save(docx_path)
        print(f"✅ ATS-optimized DOCX complete")
        return True
        
    except Exception as e:
        print(f"❌ Formatting failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Main function"""
    default_docx = "resume/main/Resume_Main.docx"
    docx_path = sys.argv[1] if len(sys.argv) > 1 else default_docx
    
    if not os.path.exists(docx_path):
        print(f"❌ File not found: {docx_path}")
        sys.exit(1)
    
    print(f"🚀 Creating ATS-optimized DOCX: {docx_path}")
    
    if compress_docx_formatting(docx_path):
        print(f"✅ DOCX compression completed successfully!")
        print(f"🏆 Professional 1-page DOCX ready")
        print(f"🎉 Features: Labeled contact, single bullets, blue headers")
        print(f"📍 Location: {docx_path}")
        return True
    else:
        print(f"⚠️ Formatting incomplete")
        return False

if __name__ == "__main__":
    main()
